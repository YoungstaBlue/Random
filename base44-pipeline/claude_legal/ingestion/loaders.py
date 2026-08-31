"""Load raw legal documents from disk into RawDocument payloads.

Plain-text and markdown are read directly. PDF/DOCX are read when the optional
libraries (pypdf / python-docx) are installed; otherwise those files are skipped
with a logged warning rather than crashing ingestion.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Union

from ..schemas import RawDocument

log = logging.getLogger("claude_legal.ingestion")

_TEXT_SUFFIXES = {".txt", ".md", ".text", ".log"}


def load_text(path: Union[str, Path]) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return p.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _read_pdf(p)
    if suffix in {".docx", ".doc"}:
        return _read_docx(p)
    # Unknown extension: best-effort text read.
    return p.read_text(encoding="utf-8", errors="ignore")


def load_documents(directory: Union[str, Path]) -> List[RawDocument]:
    """Recursively load every supported document under ``directory``."""
    root = Path(directory)
    docs: List[RawDocument] = []
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        try:
            text = load_text(path)
        except Exception as exc:  # pragma: no cover - I/O edge
            log.warning("skipping %s: %s", path, exc)
            continue
        if not text.strip():
            continue
        docs.append(
            RawDocument(
                doc_id=str(path.relative_to(root)),
                text=text,
                source=str(path),
                metadata={"suffix": path.suffix.lower()},
            )
        )
    return docs


def _read_pdf(path: Path) -> str:  # pragma: no cover - optional dependency
    try:
        from pypdf import PdfReader
    except Exception:
        log.warning("pypdf not installed; skipping PDF %s", path)
        return ""
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:  # pragma: no cover - optional dependency
    try:
        import docx
    except Exception:
        log.warning("python-docx not installed; skipping DOCX %s", path)
        return ""
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)
